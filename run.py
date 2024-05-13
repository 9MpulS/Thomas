import flet as ft
import datetime
from docx import Document
import mysql.connector


def main(page: ft.Page):
    # Початкові налаштування сторінки
    page.title = "Thomas"
    page.theme_mode = "dark"
    page.window_width = 1068
    page.window_height = 720
    page.window_resizable = False

    # Функція регестрації
    def register(e):
        # Підключення до бд
        try:
            db = mysql.connector.connect(
                host="localhost",
                user="root",
                passwd="@Vlad1234",
                database="thomas"
            )
            cursor = db.cursor()

            # Перевірка на існування логіну в таблиці users
            check_sql = """SELECT * FROM users WHERE login = %s"""
            check_val = (user_login.value,)
            cursor.execute(check_sql, check_val)
            if cursor.fetchone():
                page.snack_bar = ft.SnackBar(ft.Text("Користувач з таким логіном вже існує"))
                page.snack_bar.open = True
                page.update()
                db.close()
                return

            # Додавання даних до таблиці users
            sql = """INSERT INTO users (login, passwd) VALUES (%s, %s)"""
            val = (user_login.value, user_passwd.value)
            cursor.execute(sql, val)
            db.commit()
            db.close()

            # Очищення полів
            user_login.value = ""
            user_passwd.value = ""
            btn_reg.disabled = True
            page.update()

        except Exception as ex:
            print("Error:", ex)

    # Функція для активації кнопок
    def validate(e):
        if all([user_login.value, user_passwd.value]):
            btn_reg.disabled = False
            btn_auth.disabled = False
        else:
            btn_reg.disabled = True
            btn_auth.disabled = True
        page.update()

    # Функція для авторизації
    def authorizate(e):
        try:
            db = mysql.connector.connect(
                host="localhost",
                user="root",
                passwd="@Vlad1234",
                database="thomas"
            )
            cursor = db.cursor()

            # Перевірка чи є введені дані в таблиці users
            sql = """SELECT * FROM users WHERE login = %s AND passwd = %s"""
            val = (user_login.value, user_passwd.value)
            cursor.execute(sql, val)
            user_data = cursor.fetchone()

            if user_data:
                user_login.value = ""
                user_passwd.value = ""
                btn_auth.text = "Авторизовано"
                page.update()
                page.clean()

                # Перевірка на адміна та перехід на наступну сторінку з відровідними правами
                if user_data[3] == 1:
                    page.add(panel_srch)
                    page.add(admin_bar)
                else:
                    page.add(panel_srch)
                    page.add(user_bar)
                page.update()

            else:
                user_login.value = ""
                user_passwd.value = ""
                page.snack_bar = ft.SnackBar(ft.Text("Невірно введені пароль або логін"))
                page.snack_bar.open = True
                page.update()
            db.close()

        except Exception as ex:
            print("Error:", ex)

    # Функція для зміни теми
    def theme_toggle(e):
        if page.theme_mode == "dark":
            page.theme_mode = "light"
            btn_theme.icon = ft.icons.LIGHT_MODE_OUTLINED
        else:
            page.theme_mode = "dark"
            btn_theme.icon = ft.icons.BRIGHTNESS_3
        page.update()

    # Функція для занесення обраних потягів в файл
    def download_pick(e):
        doc = Document()
        doc.add_heading('Таблиця потягів', level=1)
        # table_headers = [column.label.value for column in table.columns]
        # doc.add_paragraph('\t'.join(table_headers))
        for row in table.rows:
            row_data = [cell.content.value for cell in row.cells]
            doc.add_paragraph('\t'.join(row_data))
        doc.save('train_table.docx')

    # Функція для очистки текстових полів
    def clean_pick(e):
        num_train.value = ""
        end_point.value = ""
        date_text.value = ""
        time_text.value = ""
        trvl_time.value = ""
        places.value = ""
        page.update()
    # Функція для обробки подій подвійного тапу по рядках таблиці
    def set_row(e):
        selected_row = e.control
        if selected_row:
            num_train.value = selected_row.cells[0].content.value
            end_point.value = selected_row.cells[1].content.value
            date_text.value = selected_row.cells[2].content.value
            time_text.value = selected_row.cells[3].content.value
            trvl_time.value = selected_row.cells[4].content.value
            places.value = selected_row.cells[5].content.value
            page.update()


    # Таблиця для запису результатів пошуку
    table = ft.DataTable(
        border=ft.border.all(1),
        columns=[
            ft.DataColumn(ft.Text("Номер потяга")),
            ft.DataColumn(ft.Text("Місце призначення")),
            ft.DataColumn(ft.Text("Дата відправлення")),
            ft.DataColumn(ft.Text("Час відправлення")),
            ft.DataColumn(ft.Text("Час в дорозі")),
            ft.DataColumn(ft.Text("К-ть місць")),
        ]
    )

    # Функція для пашуку в бд потягів
    def srch_train(e):
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="@Vlad1234",
            database="thomas"
        )
        cursor = db.cursor()

        # Пошук потягів в таблиці trains
        sql = "SELECT * FROM trains"
        val = ()
        if end_point.value:
            sql += " WHERE end_point = %s"
            val += (end_point.value,)
        if date_text.value:
            if end_point.value:
                sql += " AND "
            else:
                sql += " WHERE "
            sql += " strt_date = %s"
            val += (date_text.value,)
        if time_text.value:
            if end_point.value or date_text.value:
                sql += " AND "
            else:
                sql += " WHERE "
            sql += " strt_time >= %s"
            val += (time_text.value,)
        cursor.execute(sql, val)
        trains = cursor.fetchall()

        # Видалення таблиці якщо вона є на сторінці
        table.rows.clear()
        try:
            page.remove_at(2)
        except:
            pass

        # Стоврення таблиці
        if trains:
            for train in trains:
                t_row = ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Text(str(train[0]))),  # Номер потяга
                        ft.DataCell(ft.Text(str(train[1]))),  # Місце призначення
                        ft.DataCell(ft.Text(str(train[2]))),  # Дата відправлення
                        ft.DataCell(ft.Text(str(train[3]))),  # Час відправлення
                        ft.DataCell(ft.Text(str(train[4]))),  # Час в дорозі
                        ft.DataCell(ft.Text(str(train[5]))),  # К-ть місць
                    ], on_select_changed=set_row
                )
                table.rows.append(t_row)
            t_bar = ft.Row(
                [
                    table
                ], alignment=ft.MainAxisAlignment.CENTER
            )
            page.add(t_bar)
        else:
            page.snack_bar = ft.SnackBar(ft.Text("За цими параметрами потягів не знайдено"))
            page.snack_bar.open = True
        page.scroll = "always"
        page.update()
        db.close()

    # Функція для редагування даних
    def edit_info(e):
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="@Vlad1234",
            database="thomas"
        )
        cursor = db.cursor()

        sql = "UPDATE trains SET end_point = %s, strt_date = %s, strt_time = %s, trvl_time = %s, places = %s WHERE num = %s"
        val = (end_point.value, date_text.value, time_text.value, trvl_time.value, places.value, num_train.value)
        cursor.execute(sql, val)
        db.commit()
        db.close()
        srch_train(e)
        page.update()

    # Функція для додавання нових даних
    def add_info(e):
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="@Vlad1234",
            database="thomas"
        )
        cursor = db.cursor()

        sql = "INSERT INTO trains (num, end_point, strt_date, strt_time, trvl_time, places) VALUES (%s, %s, %s, %s, %s, %s)"
        val = (num_train.value, end_point.value, date_text.value, time_text.value, trvl_time.value, places.value)
        cursor.execute(sql, val)
        db.commit()
        db.close()
        srch_train(e)
        page.update()

    # Функція для видалення даних
    def del_info(e):
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="@Vlad1234",
            database="thomas"
        )
        cursor = db.cursor()
        sql = "DELETE FROM trains WHERE num = %s"
        val = (num_train.value,)
        cursor.execute(sql, val)
        db.commit()
        db.close()
        srch_train(e)
        page.update()

    # Кнопки та текстові поля для авторизації/регестрації
    user_login = ft.TextField(label="Логін", width=200, on_change=validate)
    user_passwd = ft.TextField(label="Пароль", width=200, password=True, can_reveal_password=True, on_change=validate)
    btn_reg = ft.OutlinedButton(text="Зареєструватися", width=200, on_click=register, disabled=True)
    btn_auth = ft.OutlinedButton(text="Увійти", width=200, on_click=authorizate, disabled=True)

    # Кнопка зміни теми
    btn_theme = ft.IconButton(icon=ft.icons.BRIGHTNESS_3, on_click=theme_toggle)

    # Панель теми
    top_bar = ft.Row(
        [
            btn_theme,
        ], alignment=ft.MainAxisAlignment.START
    )

    # Панель для регестрації
    panel_reg = ft.Row(
        [
            ft.Column(
                [
                    ft.Text("Реєстрація"),
                    user_login,
                    user_passwd,
                    btn_reg
                ], alignment=ft.MainAxisAlignment.CENTER
            )
        ], alignment=ft.MainAxisAlignment.CENTER
    )

    # Панель для авторизації
    panel_auth = ft.Row(
        [
            ft.Column(
                [
                    ft.Text("Авторизація"),
                    user_login,
                    user_passwd,
                    btn_auth
                ]
            )
        ], alignment=ft.MainAxisAlignment.CENTER
    )

    # Текстові поля для дати та часу відправлення
    date_text = ft.TextField(label="Дата відправлення", width=200, hint_text="YYYY/MM/DD")
    time_text = ft.TextField(label="Час відправлення", width=200, hint_text="hh:mm")

    # Функції для зміни дати та часу
    def change_date(e):
        date_text.value = date_picker.value.strftime('%Y/%m/%d')
        page.update()

    def change_time(e):
        time_text.value = time_picker.value.strftime('%H:%M')
        page.update()

    # Вікна для вибору дати та часу
    date_picker = ft.DatePicker(
        on_change=change_date,
        first_date=datetime.datetime(2023, 1, 1),
        last_date=datetime.datetime(2026, 1, 1),
    )
    time_picker = ft.TimePicker(
        confirm_text="Confirm",
        error_invalid_text="Time out of range",
        help_text="Pick your time slot",
        on_change=change_time,
    )
    page.overlay.append(date_picker)
    page.overlay.append(time_picker)

    # Текстові поля та кнопки для пошуку/редагування інформації про потягів
    end_point = ft.TextField(label="Станція призначення", width=400, on_change=validate)
    strt_date = ft.IconButton(icon=ft.icons.DATE_RANGE, on_click=lambda _: date_picker.pick_date())
    strt_time = ft.IconButton(icon=ft.icons.ACCESS_TIME, on_click=lambda _: time_picker.pick_time())

    btn_dwnld = ft.IconButton(icon=ft.icons.FILE_DOWNLOAD_OUTLINED, on_click=download_pick)
    btn_clean = ft.IconButton(icon=ft.icons.RESTART_ALT, on_click=clean_pick)
    btn_srhc = ft.OutlinedButton(text="Пошук", width=200, on_click=srch_train)

    num_train = ft.TextField(label="Номер потяга", width=100)
    trvl_time = ft.TextField(label="Час в дорозі", width=100)
    places = ft.TextField(label="К-ть місць", width=100)

    btn_edit = ft.OutlinedButton(text="Зберегти зміни", width=200, on_click=edit_info)
    btn_add = ft.OutlinedButton(text="Додати", width=200, on_click=add_info)
    btn_del = ft.OutlinedButton(text="Видалити", width=200, on_click=del_info)

    # Панель для пошуку
    panel_srch = ft.Column(
        [
            ft.Row(
                [
                    btn_theme
                ], alignment=ft.MainAxisAlignment.START
            ),
            ft.Row(
                [
                    end_point,
                ], alignment=ft.MainAxisAlignment.CENTER
            ),
            ft.Row(
                [
                    strt_date,
                    date_text,
                    time_text,
                    strt_time
                ], alignment=ft.MainAxisAlignment.CENTER
            ),
            ft.Row(
                [
                    btn_dwnld,
                    btn_srhc,
                    btn_clean
                ], alignment=ft.MainAxisAlignment.CENTER
            )
        ], alignment=ft.MainAxisAlignment.CENTER
    )

    # Панель для редагування
    panel_edit = ft.Column(
        [
            ft.Row(
                [
                    btn_theme
                ], alignment=ft.MainAxisAlignment.START
            ),
            ft.Row(
                [
                    num_train,
                    end_point,
                    places
                ], alignment=ft.MainAxisAlignment.CENTER
            ),
            ft.Row(
                [
                    strt_date,
                    date_text,
                    time_text,
                    strt_time,
                    trvl_time
                ], alignment=ft.MainAxisAlignment.CENTER
            ),
            ft.Row(
                [
                    btn_dwnld,
                    btn_srhc,
                    btn_clean
                ], alignment=ft.MainAxisAlignment.CENTER
            ),
            ft.Row(
                [
                    btn_add,
                    btn_edit,
                    btn_del
                ], alignment=ft.MainAxisAlignment.CENTER
            )
        ], alignment=ft.MainAxisAlignment.CENTER
    )

    # Функція для навігації по стоінкі авторизації
    def navigate_outh(e):
        index = auth_bar.selected_index
        page.clean()
        if index == 0:
            page.add(top_bar)
            page.add(panel_reg)
            page.add(auth_bar)
        elif index == 1:
            page.add(top_bar)
            page.add(panel_auth)
            page.add(auth_bar)

    # Функція для нафігації з правами user
    def navigate_user(e):
        index = user_bar.selected_index
        page.clean()
        if index == 0:
            page.add(panel_srch)
            page.add(user_bar)
        elif index == 1:
            page.add(top_bar)
            page.add(panel_auth)
            page.add(auth_bar)

    # Функція для нафігації з правами admin
    def navigate_admin(e):
        index = admin_bar.selected_index
        page.clean()
        if index == 0:
            page.add(panel_srch)
            page.add(admin_bar)
        elif index == 1:
            page.add(panel_edit)
            page.add(admin_bar)
        elif index == 2:
            page.add(top_bar)
            page.add(panel_auth)
            page.add(auth_bar)

    # Навігаційна панель сторінки авторизації
    auth_bar = ft.NavigationBar(
        destinations=[
            ft.NavigationDestination(icon=ft.icons.HOW_TO_REG_OUTLINED, label="Реєстрація"),
            ft.NavigationDestination(icon=ft.icons.HOW_TO_REG, label="Авторизація")
        ], on_change=navigate_outh
    )

    # Навігаційна панель основної сторінки для user
    user_bar = ft.NavigationBar(
        destinations=[
            ft.NavigationDestination(icon=ft.icons.SEARCH, label="Пошук квитків"),
            ft.NavigationDestination(icon=ft.icons.EXIT_TO_APP, label="Вихід")
        ], on_change=navigate_user
    )

    # Навігаційна панель основної сторінки для admin
    admin_bar = ft.NavigationBar(
        destinations=[
            ft.NavigationDestination(icon=ft.icons.SEARCH, label="Пошук квитків"),
            ft.NavigationDestination(icon=ft.icons.EDIT, label="Редактор"),
            ft.NavigationDestination(icon=ft.icons.EXIT_TO_APP, label="Вихід")
        ], on_change=navigate_admin
    )

    # Стартові налаштунки сторінки
    page.add(top_bar)
    page.add(panel_reg)
    page.add(auth_bar)


# Запуск
ft.app(target=main)
